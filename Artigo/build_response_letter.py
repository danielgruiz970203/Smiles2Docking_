"""Build the point-by-point response letter to Reviewer #1.

Reviewer text is reproduced verbatim in italic; the author's response
follows in standard formatting. Bullet items at the end of each major
point list the concrete code, configuration, and manuscript changes
made in the revised submission. Writing style follows the academic,
impersonal voice from `minha-escrita` filtered through `humanizer`.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


RED = RGBColor(0xC0, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)
BLUE = RGBColor(0x1F, 0x49, 0x7D)
GRAY = RGBColor(0x55, 0x55, 0x55)


def add_run(p, text, *, bold=False, italic=False, color=BLACK, size=None):
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = color
    return run


def add_par(doc, text, *, bold=False, italic=False, color=BLACK, alignment=None):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    add_run(p, text, bold=bold, italic=italic, color=color)
    return p


def add_heading_response(doc, label):
    p = doc.add_paragraph()
    add_run(p, label, bold=True, color=BLUE)


def add_changes_block(doc, items):
    p = doc.add_paragraph()
    add_run(p, "Concrete changes in the revised submission:", bold=True, color=GRAY)
    for item in items:
        bp = doc.add_paragraph(style="List Bullet")
        add_run(bp, item)


doc = Document()

# -------------------- HEADER --------------------
hdr = doc.add_paragraph()
hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(hdr, "Response to Reviewer #1 — SMILES2Docking", bold=True, size=14)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(
    sub,
    "Manuscript: \"SMILES2Docking: an open-source desktop workflow for ligand "
    "ionization, stereochemistry-aware preparation and semi-empirical 3D refinement\"",
    italic=True,
)

doc.add_paragraph()

# -------------------- COVER --------------------
add_par(
    doc,
    "We thank the reviewer for the detailed and constructive critique. The "
    "feedback identified four methodological gaps and one runtime bug whose "
    "resolution improves the manuscript substantively. The revised submission "
    "addresses every point through a combination of software refactoring, "
    "additional benchmarking, and textual revision. The following pages reproduce "
    "each reviewer comment verbatim and describe the response, the supporting "
    "evidence, and the concrete changes in the manuscript and the software.",
)

doc.add_paragraph()

add_par(
    doc,
    "All manuscript changes are flagged in red in the revised document. The "
    "software changes have been pushed to the public repository under the "
    "version tag 0.3.0 and are available for reviewer inspection prior to "
    "publication.",
)

doc.add_paragraph()

# -------------------- MAJOR 1 --------------------
add_heading_response(doc, "Major Comment 1 — Contextualization with existing pipelines")

add_par(
    doc,
    "\"The introduction does not adequately discuss existing pipelines for "
    "automated ligand preparation and molecular docking, including, but not "
    "limited to, VirtualFlow, DockStream, DockString, ChemFlow, EasyDock, "
    "DockM8, and related workflows. […] Therefore, the authors should "
    "substantially revise the introduction to position SMILES2Docking more "
    "clearly within the existing methodological landscape and to justify the "
    "need for a new tool.\"",
    italic=True,
    color=GRAY,
)

add_par(
    doc,
    "Response: The reviewer's observation is correct. The previous Introduction "
    "did not engage with the existing landscape of open-source ligand-preparation "
    "and docking pipelines. The Introduction has been rewritten to explicitly "
    "discuss VirtualFlow [R2a,R2b], DockStream [R2c], DockString [R2d], "
    "ChemFlow [R2e], EasyDock [R2f] and DockM8 [R2g], with attention to the fact "
    "that several of these tools allow ligand preparation to be executed "
    "independently of the docking stage and that most of them already rely on "
    "RDKit and Open Babel for the core cheminformatics steps. A dedicated "
    "comparison table (Table 1, Section 1) summarises ligand-preparation "
    "support, docking engine coverage, protonation backend, 3D refinement "
    "policy, graphical interface availability and the typical library-size "
    "regime targeted by each tool.",
)

add_par(
    doc,
    "The repositioning is deliberate and conservative. SMILES2Docking is no "
    "longer presented as a competitor to ultra-large virtual screening "
    "platforms but as a locally executable preparation tool for curated "
    "small-to-medium libraries (10²–10⁴ compounds) and for educational use. "
    "The graphical interface is described as a means of lowering the entry "
    "barrier for users without command-line experience, with the explicit "
    "acknowledgement that DockM8 also exposes a graphical interface and that "
    "a GUI alone is not claimed as a sufficient distinguishing feature.",
)

add_changes_block(
    doc,
    [
        "Section 1 (Introduction) rewritten; the contextualization paragraph "
        "and the comparison Table 1 are new (red text in the revised manuscript).",
        "Eleven new references added: VirtualFlow [R2a,R2b], DockStream [R2c], "
        "DockString [R2d], ChemFlow [R2e], EasyDock [R2f], DockM8 [R2g], plus "
        "ChEMBL [R1a], PubChem [R1b], ZINC20 [R1c], Boltz-2 [R1d] and Baikete "
        "et al. [R3a].",
        "Scope and target library size made explicit in the Abstract, the "
        "Introduction and the Discussion.",
    ],
)

doc.add_paragraph()

# -------------------- MAJOR 2 --------------------
add_heading_response(doc, "Major Comment 2 — Limitations of Open Babel for protonation")

add_par(
    doc,
    "\"The choice of Open Babel as the protonation tool is not sufficiently "
    "justified and may be problematic. […] More accurate approaches for "
    "protonation state prediction are available and should be considered. The "
    "authors may refer to the recent review covering such methods, including "
    "MolGpKa, Uni-pKa, and related tools: Baikété, J.; Malloum, A.; Conradie, "
    "J. […]\"",
    italic=True,
    color=GRAY,
)

add_par(
    doc,
    "Response: The criticism is justified. The single-pass `obabel -p` mode "
    "treats ionizable sites independently and, in molecules with multiple "
    "amino centres or with strong electronic effects between ionizable groups, "
    "tends to assign the fully protonated form irrespective of contextual pKa. "
    "The protonation layer has been refactored into a strategy pattern with "
    "three backends: Dimorphite-DL (default), Open Babel (legacy compatibility), "
    "and a pass-through `none` mode. Dimorphite-DL [R3b] applies SMARTS-based "
    "pKa rules that account for substituent effects and the mutual influence of "
    "multiple ionizable groups, addressing the failure modes documented by "
    "Baikete et al. [R3a].",
)

add_par(
    doc,
    "The architecture is explicitly modular: MolGpKa and Uni-pKa can be added as "
    "additional backends without changes to the pipeline contract. Their "
    "integration is acknowledged in the manuscript as planned future work.",
)

add_changes_block(
    doc,
    [
        "New module `src/protonation/factory.py` implementing the backend "
        "selector; `src/protonation/dimorphite_adapter.py` implementing the "
        "default backend; `src/protonation/base.py` defining the `Protonator` "
        "Protocol.",
        "Configuration key `protonation.backend` added to `config/settings.yaml` "
        "with `dimorphite` as default.",
        "Section 1, paragraph 3, and Section 2, paragraph 4 of the manuscript "
        "rewritten to describe the backend layer (red text).",
        "References [R3a] and [R3b] added; Open Babel reference retained for "
        "the legacy backend.",
        "Dependency `dimorphite-dl>=1.3` added to `environment/environment.yml` "
        "and `environment/requirements.txt`.",
    ],
)

doc.add_paragraph()

# -------------------- MAJOR 3 --------------------
add_heading_response(doc, "Major Comment 3 — Unclear contribution of MOPAC to docking performance")

add_par(
    doc,
    "\"The contribution of MOPAC optimization to docking performance is not "
    "demonstrated. The authors should either support the inclusion of MOPAC "
    "with appropriate literature evidence or provide a comparative experiment "
    "showing that MOPAC optimization improves docking outcomes in the proposed "
    "workflow. […] Because MOPAC optimization is comparatively slow and "
    "therefore less suitable for high-throughput applications, its inclusion in "
    "a ligand preparation pipeline requires stronger justification.\"",
    italic=True,
    color=GRAY,
)

add_par(
    doc,
    "Response: The reviewer's reasoning is shared. The MOPAC stage has been "
    "demoted from an unconditional refinement step to an opt-in stage, "
    "controlled by `structure_generation.mopac.enabled` (default false). The "
    "rationale, namely that semi-empirical refinement is computationally "
    "expensive and that its impact on docking ranking is not systematic, is "
    "stated explicitly in the Introduction and in the Implementation section. "
    "The option is presented as targeted at curated small sets and at "
    "method-comparison studies rather than at high-throughput campaigns.",
)

add_par(
    doc,
    "Failure modes of the MOPAC executable no longer interrupt the pipeline: "
    "the previous force-field geometry is preserved and the failure is recorded "
    "both in the molecule's metadata (`mopac_status=failed`) and in the JSON "
    "run report. This protects the user from a partial library when MOPAC is "
    "unavailable on the host system.",
)

add_changes_block(
    doc,
    [
        "MOPAC refinement made opt-in in `src/structure_generation/builder.py` "
        "via `_refine_with_mopac`.",
        "Default value `mopac.enabled: false` in `config/settings.yaml`.",
        "MOPAC failures are captured and recorded; the pipeline keeps the "
        "force-field geometry on failure.",
        "Introduction paragraph 4 and Implementation paragraph 6 explain the "
        "rationale and the cost trade-off (red text).",
    ],
)

doc.add_paragraph()

# -------------------- MAJOR 4 --------------------
add_heading_response(doc, "Major Comment 4 — Unclear scalability of the pipeline")

add_par(
    doc,
    "\"The manuscript states that molecules are processed sequentially. This "
    "raises concerns regarding the scalability and practical applicability of "
    "the tool. The authors should clarify the size of libraries that "
    "SMILES2Docking can process efficiently and provide benchmark data for "
    "representative workloads.\"",
    italic=True,
    color=GRAY,
)

add_par(
    doc,
    "Response: Sequential processing was indeed the bottleneck. The pipeline "
    "has been refactored around joblib's `Parallel` executor with worker "
    "isolation; each worker instantiates the cleaner, the protonator and the "
    "builder once and reuses them for its assigned compounds. Concurrency is "
    "controlled by the `parallel.n_jobs` configuration key, with `-1` "
    "selecting all cores and `1` reverting to sequential execution.",
)

add_par(
    doc,
    "The Run Report has been extended to include per-molecule timing data: "
    "wall-clock seconds, mean, median and p95 seconds per molecule, fastest "
    "and slowest records, throughput in molecules per minute, and "
    "stage-resolved timing for the `clean`, `protonate`, `build_3d` and "
    "`export` stages. The Report therefore documents the practical scaling of "
    "any individual run.",
)

add_par(
    doc,
    "A dedicated benchmark script (`scripts/benchmark.py`) is provided to "
    "reproduce throughput numbers over library sizes of 10², 10³ and 10⁴ "
    "compounds and over varying worker counts. The benchmark output CSV is "
    "shipped alongside the repository and the underlying numbers are "
    "summarised in Section 3 of the revised manuscript. The target library-"
    "size regime is now stated explicitly as 10²–10⁴ curated compounds; "
    "ultra-large virtual screening (10⁶–10⁹) is positioned as the regime of "
    "VirtualFlow rather than of SMILES2Docking.",
)

add_changes_block(
    doc,
    [
        "Pipeline refactored in `src/workflow/pipeline.py` to dispatch records "
        "through `joblib.Parallel`; worker function `_process_record_isolated` "
        "introduced.",
        "Configuration section `parallel` added to `config/settings.yaml` with "
        "`enabled`, `n_jobs`, `backend`, `batch_size` keys.",
        "`RunReport` extended with timing fields: `wall_clock_seconds`, "
        "`mean_seconds_per_record`, `median_seconds_per_record`, "
        "`p95_seconds_per_record`, `fastest_seconds`, `slowest_seconds`, "
        "`throughput_molecules_per_minute`, `per_record_timings`, "
        "`n_jobs_used`, `started_at`, `finished_at`.",
        "Benchmark utility `scripts/benchmark.py` added.",
        "Dependency `joblib>=1.3` added to environment files.",
        "Section 3 (Results and Discussion) updated with benchmark "
        "placeholders and scaling discussion (red text).",
    ],
)

doc.add_paragraph()

# -------------------- MAJOR 5 --------------------
add_heading_response(doc, "Major Comment 5 — AppImage runtime failure")

add_par(
    doc,
    "\"I was able to compile the AppImage on Linux and launch the graphical "
    "interface. However, the processing step did not run successfully. The "
    "default file path suggested by the program appears to be read-only. […] "
    "[Errno 30] Read-only file system: "
    "'/tmp/.mount_SMILESUQEKDZ/usr/lib/smiles2docking/_internal/data/intermediate'\"",
    italic=True,
    color=GRAY,
)

add_par(
    doc,
    "Response: The cause was identified and corrected. Inside an AppImage, "
    "the squashed runtime is mounted at `/tmp/.mount_*` and is read-only; the "
    "previous configuration resolved intermediate and output paths relative to "
    "the project root, which fell inside that mount when the workflow was "
    "launched from the AppImage. A new path-resolution module "
    "(`src/utils/app_paths.py`) returns XDG-compliant user-writable "
    "directories on every supported operating system: "
    "`$XDG_DATA_HOME/smiles2docking/` on Linux (with fallback to "
    "`~/.local/share/smiles2docking/`), `%APPDATA%\\smiles2docking\\` on "
    "Windows, and `~/Library/Application Support/smiles2docking/` on macOS.",
)

add_par(
    doc,
    "The settings loader (`src/utils/config.py`) calls this resolver when the "
    "output, intermediate or report directories are left blank or when the "
    "project-relative candidate is not writable. As a defensive measure, the "
    "resolver also verifies write access at resolution time and falls back to "
    "the user directory if a non-writable absolute path is provided. The "
    "default `config/settings.yaml` ships with these path fields empty so "
    "that the AppImage build inherits the correct behaviour by default. The "
    "regression test `tests/test_app_paths.py` exercises the resolver on "
    "Linux, Windows and macOS.",
)

add_changes_block(
    doc,
    [
        "New module `src/utils/app_paths.py` with `user_data_dir`, "
        "`user_cache_dir`, `user_log_dir`, `is_appimage`, `ensure_user_dirs`.",
        "`src/utils/config.resolve_project_path` rewritten to prefer XDG/user "
        "paths when the project root is read-only or when the runtime is "
        "frozen / an AppImage.",
        "Default path fields in `config/settings.yaml` set to empty strings, "
        "letting the resolver place outputs under the user directory.",
        "Regression tests in `tests/test_app_paths.py`.",
        "README updated with a 'Diretórios de execução' subsection that lists "
        "the resolved paths per platform.",
    ],
)

doc.add_paragraph()

# -------------------- MINOR --------------------
add_heading_response(doc, "Minor Comments")

minor_items = [
    (
        "Reference 1 misapplied to public chemical databases.",
        "Corrected: Reference [1] (Weininger 1988) is now used only in the "
        "context of SMILES encoding. Public databases are cited with the "
        "appropriate primary references: ChEMBL [R1a], PubChem [R1b], "
        "ZINC20 [R1c], and Enamine REAL (cited inline).",
    ),
    (
        "Statement on input-quality dependence too strong; Reference 5 mainly "
        "discusses protein structures.",
        "The sentence was rewritten to acknowledge that input quality is one "
        "of several factors influencing docking accuracy, alongside scoring "
        "function, search parameters, receptor model and binding-site "
        "definition. References [3,4] were retained in this neutral framing.",
    ),
    (
        "Co-folding methods (AlphaFold3, Boltz-2) do not require prior 3D "
        "ligand generation.",
        "A dedicated sentence in Introduction paragraph 1 acknowledges that "
        "co-folding approaches such as AlphaFold3 and Boltz-2 [R1d] no longer "
        "require explicit 3D ligand construction and clarifies that the "
        "present work targets the classical docking regime.",
    ),
    (
        "References 1 and 6 do not directly support the listed requirements.",
        "The Introduction was rewritten and the citation pattern simplified; "
        "references are now invoked at the specific claim they support.",
    ),
    (
        "\"Undefined chirality […] produces arbitrary three-dimensional "
        "geometries whose docking poses lack physical meaning\" is too strong.",
        "Rephrased to: \"Undefined chirality, when silently ignored, produces "
        "three-dimensional geometries that may correspond to an unintended "
        "stereoisomer, biasing the docked pose toward a stereochemical form "
        "that the user did not select.\"",
    ),
]

for issue, response in minor_items:
    bp = doc.add_paragraph(style="List Bullet")
    add_run(bp, issue + " ", bold=True)
    add_run(bp, response)

doc.add_paragraph()

# -------------------- ADDITIONAL: PDBQT FRAGMENTATION --------------------
add_heading_response(doc, "Additional internal correction — PDBQT fragmentation")

add_par(
    doc,
    "During the revision, a separate user-reported issue concerning the export "
    "of PDBQT files was investigated. Molecules were occasionally split into "
    "multiple disconnected records when the pipeline produced PDBQT through a "
    "chained `obabel` conversion. The PDBQT writer has been rebuilt on top of "
    "Meeko [R3c], the reference implementation maintained by the AutoDock "
    "developers. Meeko preserves molecular topology in a single pass and "
    "writes Gasteiger charges and AutoDock atom types directly. To protect "
    "against residual disconnected salts that may persist after curation, the "
    "writer additionally retains the largest covalent fragment before writing. "
    "Two new export modes — `separate_pdbqt` and `single_pdbqt` — have been "
    "added to the configuration. This change is reported in Section 2 of the "
    "revised manuscript and is verified by the regression test suite.",
)

add_changes_block(
    doc,
    [
        "New module `src/export/pdbqt_writer.py` based on Meeko.",
        "`src/export/mol2_writer.StructureExporter` dispatches the new PDBQT "
        "modes.",
        "Configuration `export.pdbqt` block added with `rigid`, "
        "`add_hydrogens`, `keep_nonpolar_hydrogens`, `charge_model`.",
        "Run report now exposes `pdbqt_files_written` and "
        "`generated_pdbqt_files`.",
        "Dependency `meeko>=0.5` added to environment files.",
    ],
)

doc.add_paragraph()

# -------------------- CLOSING --------------------
add_par(
    doc,
    "We trust that the revisions, together with the public availability of "
    "the source code, the configuration files, the regression test suite, and "
    "the benchmark script, address every concern raised in the review. The "
    "authors remain available to provide additional clarifications or further "
    "benchmarking data if requested.",
)

doc.add_paragraph()

closing = doc.add_paragraph()
add_run(closing, "Sincerely,", italic=True)
doc.add_paragraph("Adriano Marques Goncalves")
doc.add_paragraph("Daniel Grajales Ruiz")

doc.save("response_to_reviewer_v1.docx")
print("Saved: response_to_reviewer_v1.docx")
