"""
Writes the humanized version of smiles2docking_article.docx.
Run once: python humanize_article.py
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── Title ──────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run(
    "SMILES2DockingFULL: an open-source desktop and command-line workflow "
    "for stereochemistry-aware ligand preparation with PM7 semi-empirical refinement"
)
run.bold = True
run.font.size = Pt(14)

# ── Authors ────────────────────────────────────────────────────────────────
authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors.add_run(
    "Adriano Marques Goncalves¹, Daniel Grajales Ruiz²\n"
    "¹ Universidade de Araraquara — UNIARA, Araraquara, SP, Brazil\n"
    "² Instituto de Química, Universidade Estadual Paulista — UNESP, Araraquara, SP, Brazil\n"
    "Correspondence: amgoncalves@uniara.edu.br"
)

doc.add_paragraph()

# ── Abstract ───────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.add_run("Abstract").bold = True

doc.add_paragraph(
    "SMILES2DockingFULL converts spreadsheet-based SMILES libraries into docking-ready 3D ligands "
    "through a single pipeline. Salt and co-former removal, a three-mode stereocentre policy, "
    "pH-dependent protonation via Open Babel, distance-geometry embedding with RDKit, molecular "
    "mechanics optimization, and optional PM7 semi-empirical refinement under implicit solvation are "
    "all handled in sequence. Users choose among permissive, strict, and enumerative stereocentre "
    "modes, allowing the same tool to serve broad exploratory screening and precision docking alike. "
    "A per-compound JSON audit report records protonation and stereocentre decisions throughout. "
    "The tool ships as a Python package, a Windows executable with bundled MOPAC, and a Linux "
    "portable bundle under GPL-2.0-or-later. SMILES2DockingFULL is freely available to "
    "non-commercial users at https://github.com/amgoncalvesusp/Smiles2Docking "
    "(DOI: [ZENODO_DOI_PLACEHOLDER])."
)

# ── Keywords ───────────────────────────────────────────────────────────────
kw = doc.add_paragraph()
kw.add_run("Keywords: ").bold = True
kw.add_run(
    "SMILES; ligand preparation; molecular docking; stereochemistry enumeration; protonation state; "
    "semi-empirical methods; PM7; RDKit; Open Babel; MOPAC."
)

doc.add_paragraph()

# ── Introduction placeholder ───────────────────────────────────────────────
p = doc.add_paragraph()
p.add_run("[NOTE TO AUTHORS: paste the merged Introduction here]").italic = True

doc.add_paragraph()

# ── Implementation ─────────────────────────────────────────────────────────
impl_paragraphs = [
    (
        "SMILES2DockingFULL runs in Python 3.11, either as a PySide6 desktop application or from the "
        "command line, with both interfaces sharing the same configuration. Input is read from CSV, XLS "
        "or XLSX spreadsheets via pandas; column names are matched case-insensitively, so input files "
        "need no reformatting. Each record carries its original row index through the entire pipeline, "
        "preserving a direct link between the spreadsheet and the exported structure."
    ),
    (
        "SMILES strings encoding multiple disconnected fragments (typical for salts and co-crystallised "
        "counter-ions) are handled by ranking fragments on organic content and heavy-atom count and "
        "keeping the top-scoring one. When two fragments score identically and conservative mode is "
        "active, the run halts with an explicit error instead of guessing. In a large batch over a "
        "public library, a silent wrong choice would propagate unnoticed into every downstream step, "
        "so the failure is deliberate."
    ),
    (
        "The stereocentre policy is one of three modes selected at run time. Permissive mode passes all "
        "molecules regardless of whether their tetrahedral stereocentres are assigned—the right "
        "default when working from databases like ChEMBL or PubChem, which contain many structures with "
        "incomplete stereo annotation. Strict mode rejects any compound with one or more undefined "
        "tetrahedral stereocentres, limiting the output to geometrically unambiguous structures. "
        "Enumerative mode generates all possible diastereomers up to a configurable cap (default 16) "
        "and tags each variant with a unique identifier derived from the parent code, so all downstream "
        "files remain individually traceable. Geometric E/Z stereocentres at double bonds are not "
        "currently evaluated."
    ),
    (
        "Open Babel assigns protonation at a user-specified pH (default 7.4). When the resulting SMILES "
        "cannot be re-parsed by RDKit, an SDF intermediate is used instead. The formal charge of the "
        "protonated species is recorded and carried into the post-refinement charge audit."
    ),
    (
        "RDKit generates 3D coordinates from a distance-geometry algorithm [14]. The workflow starts "
        "with a reproducible embedding at a fixed random seed; if that fails, it tries a random-seed "
        "variant, then a small-ring-adapted variant for strained or macrocyclic systems, with multiple "
        "attempts at each stage before moving on. The embedded geometry is then minimised through "
        "successive molecular mechanics force fields [15,16] until convergence. If no force field "
        "converges, the record is rejected unless the user has explicitly enabled export of "
        "unoptimised structures."
    ),
    (
        "When PM7 refinement is enabled, the converged geometry is submitted to MOPAC with a standard "
        "amide and peptide-bond correction. An implicit aqueous solvent model (ε = 78.4) "
        "[17] is available to better match physiological conditions. For method-comparison work, "
        "seven additional Hamiltonians are accessible: PM6, PM6-D3H4X, PM6-ORG, RM1, PM3, AM1 and "
        "MNDO. The heat of formation is stored as a molecular property. After refinement, a charge audit "
        "compares the MOPAC formal charge against the pre-refinement value; a discrepancy triggers an "
        "automated rescue heuristic, and if that fails the record is rejected rather than passed on with "
        "a corrupted protonation state. On Windows, MOPAC ships inside the executable and needs no "
        "separate installation."
    ),
    (
        "Structures are exported as MOL2 or SDF, either per compound or as a concatenated file. Every "
        "run also produces a JSON audit report with per-compound entries covering co-former removal, "
        "protonation, stereocentre outcomes, embedding and optimisation results, and semi-empirical "
        "metadata. For anyone troubleshooting an unexpected docking result, the report is the first "
        "place to look."
    ),
]

for p_text in impl_paragraphs:
    doc.add_paragraph(p_text)

doc.add_paragraph()

# ── Results and Discussion ─────────────────────────────────────────────────
rd_paragraphs = [
    (
        "We ran [PLACEHOLDER — N] compounds from [PLACEHOLDER — source library] at pH 7.4 "
        "with PM7 refinement and implicit aqueous solvation. The workflow returned "
        "[PLACEHOLDER — N1] cleaned structures; [PLACEHOLDER — N2] had at least one co-former "
        "removed by fragment ranking. [PLACEHOLDER — N3] records were excluded by the strict "
        "stereocentre policy and [PLACEHOLDER — N4] were expanded by enumerative mode into "
        "[PLACEHOLDER — N5] uniquely labelled stereoisomers. Embedding succeeded on the first pass "
        "for [PLACEHOLDER — N6] compounds; [PLACEHOLDER — N7] needed the random-seed fallback "
        "and [PLACEHOLDER — N8] the small-ring-adapted stage. PM7 completed for "
        "[PLACEHOLDER — N9] structures at [PLACEHOLDER — t] s per ligand on "
        "[PLACEHOLDER — hardware]. Mean heat-of-formation difference between the force-field and "
        "PM7 minima: [PLACEHOLDER — value] kcal/mol."
    ),
    (
        "Compared to a minimal RDKit-plus-Open-Babel pipeline, SMILES2DockingFULL adds two controls "
        "targeting known failure modes. The stereocentre policies make chirality handling explicit and "
        "per-compound auditable: standard RDKit embedding silently assigns an arbitrary configuration "
        "to an undefined stereocentre, whereas strict and enumerative modes filter or expand the "
        "compound. This matters because chirality errors propagate directly into binding-affinity "
        "predictions [7]. The post-refinement charge audit catches protonation changes that can occur "
        "during semi-empirical optimisation and either rescues the structure or rejects it, since an "
        "incorrect charge reliably degrades docking scores [8]."
    ),
    (
        "The tool is also aimed at researchers who do not have the time or programming background to "
        "assemble a custom preparation pipeline. ChEMBL, ZINC and PubChem together hold hundreds of "
        "millions of compounds as SMILES strings, and the steps between a SMILES and a docked pose are "
        "genuinely non-trivial. The three stereocentre modes reflect decisions that come up in practice: "
        "permissive mode for exploratory campaigns where the database contains many structures with "
        "undefined stereochemistry; strict mode when every output must be geometrically unambiguous; "
        "enumerative mode when the scaffold is fixed and all diastereomers need docking. On Windows, "
        "MOPAC ships in the executable, so semi-empirical refinement works out of the box."
    ),
    (
        "The current version retains only the dominant protonation state from Open Babel; tautomers are "
        "not enumerated. Processing is sequential, so throughput scales linearly with library size. "
        "Support for E/Z filtering at double bonds is not yet implemented but is planned."
    ),
]

for p_text in rd_paragraphs:
    doc.add_paragraph(p_text)

doc.add_paragraph()

# ── Availability and Requirements ─────────────────────────────────────────
p = doc.add_paragraph()
p.add_run("Availability and Requirements").bold = True

avail_fields = [
    ("Project name: ", "SMILES2DockingFULL (version 0.2.0)."),
    (
        "Repository: ",
        "https://github.com/amgoncalvesusp/Smiles2Docking (DOI: [ZENODO_DOI_PLACEHOLDER]). "
        "SMILES2DockingFULL is freely available to non-commercial users."
    ),
    (
        "Operating systems: ",
        "Windows 10/11 (signed executable bundle; the Windows executable bundles MOPAC, enabling "
        "zero-configuration PM7 refinement on that platform); Linux x86_64 (portable tar.gz and "
        "optional AppImage); macOS via the Python source distribution."
    ),
    ("Programming language: ", "Python 3.11."),
    (
        "Other requirements: ",
        "RDKit, Open Babel, pandas, openpyxl, PyYAML, PySide6, PyInstaller (build only), pytest "
        "(tests). Pinned dependency versions are specified in the repository's environment file. "
        "MOPAC (Apache-2.0 licence for current open-source releases) is required for PM7 refinement; "
        "it is bundled in the Windows executable and must be installed separately on Linux and macOS."
    ),
    (
        "Installation: ",
        "conda env create -f environment/environment.yml && conda activate smiles2docking. "
        "Launch the GUI with python scripts/run_gui.py or the CLI with "
        "python scripts/run_workflow.py --input <file> --ph 7.4 --pm7 --pm7-solvent --pm7-eps 78.39."
    ),
    ("Licence: ", "GPL-2.0-or-later, chosen for compatibility with the Open Babel runtime dependency."),
]

for label, content in avail_fields:
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    p.add_run(content)

doc.add_paragraph()

# ── Acknowledgements ───────────────────────────────────────────────────────
p = doc.add_paragraph()
p.add_run("Acknowledgements").bold = True
doc.add_paragraph(
    "[PLACEHOLDER — funding agencies, grant numbers, and institutional support to be confirmed "
    "by the authors before submission.]"
)

doc.add_paragraph()

# ── ACTION REQUIRED note ───────────────────────────────────────────────────
p = doc.add_paragraph()
p.add_run(
    "[ACTION REQUIRED: the citation numbers below are assigned on the assumption that the merged "
    "Introduction consumes references [1]–[13] in the order established during the previous "
    "drafting session (Weininger [1], Pinzi & Rastelli [2], Bhatt et al. [3], Gautam et al. [4], "
    "Neelam [5], Landrum/RDKit [6], Brooks et al. [7], ten Brink & Exner [8], O’Boyle et al. "
    "[9], Hawkins et al. [10], Sadowski & Gasteiger [11], Azam & Almahmoud [12], Stewart [13]). "
    "When the Introduction is pasted in, verify that every in-text number aligns with its intended "
    "reference and renumber the full list if citations were added or removed.]"
).italic = True

doc.add_paragraph()

# ── References ─────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.add_run("References").bold = True

references = [
    "[1] Weininger D. SMILES, a chemical language and information system. 1. Introduction to "
    "methodology and encoding rules. J Chem Inf Comput Sci. 1988;28(1):31–36.",

    "[2] Pinzi L, Rastelli G. Molecular Docking: Shifting Paradigms in Drug Discovery. Int J Mol Sci. "
    "2019;20(18):4331. doi:10.3390/ijms20184331",

    "[3] Bhatt A, Panda SK, Chaudhari SP, Pathak M, Satapathy A, Prasanna NK. Mapping of Global "
    "Research Performance on Molecular Docking: A Bibliometric Study. Curr Trends Biotechnol Pharm. "
    "2024;18(2):1725–1735. doi:10.5530/ctbp.2024.2.21",

    "[4] Gautam S, Pathak S, Dubey SH. The Role of Molecular Docking in Modern Drug Discovery and "
    "Development: A Comprehensive Review. J Drug Discov Health Sci. 2024;1(3):129–137. "
    "doi:10.21590/jddhs.01.03.02",

    "[5] Neelam AK. Advancing drug discovery: the role of computer-aided design and development in "
    "modern pharmaceuticals. Discov Pharm Sci. 2025;1(1):8. doi:10.1007/s44395-025-00008-2",

    "[6] Landrum G. RDKit: Open-source cheminformatics software. https://www.rdkit.org (accessed 2026).",

    "[7] Brooks WH, Guida WC, Daniel KG. The significance of chirality in drug design and development. "
    "Curr Top Med Chem. 2011;11(7):760–770.",

    "[8] ten Brink T, Exner TE. Influence of protonation, tautomeric, and stereoisomeric states on "
    "protein-ligand docking results. J Chem Inf Model. 2009;49(6):1535–1546.",

    "[9] O’Boyle NM, Banck M, James CA, Morley C, Vandermeersch T, Hutchison GR. Open Babel: "
    "an open chemical toolbox. J Cheminform. 2011;3:33.",

    "[10] Hawkins PCD, Skillman AG, Warren GL, Ellingson BA, Stahl MT. Conformer generation with "
    "OMEGA: algorithm and validation using high-quality structures from the Protein Data Bank and "
    "Cambridge Structural Database. J Chem Inf Model. 2010;50(4):572–584.",

    "[11] Sadowski J, Gasteiger J. From atoms and bonds to three-dimensional atomic coordinates: "
    "automatic model builders. Chem Rev. 1993;93(7):2567–2581.",

    "[12] Azam F, Almahmoud SA. Open-Source Molecular Docking and AI-Augmented Structure-Based Drug "
    "Design: Current Workflows, Challenges, and Opportunities. Int J Mol Sci. 2026;27(7):3302. "
    "doi:10.3390/ijms27073302",

    "[13] Stewart JJP. Optimization of parameters for semiempirical methods VI: more modifications to "
    "the NDDO approximations and re-optimization of parameters. J Mol Model. 2013;19(1):1–32.",

    "[14] Riniker S, Landrum GA. Better informed distance geometry: using what we know to improve "
    "conformation generation. J Chem Inf Model. 2015;55(12):2562–2574. "
    "doi:10.1021/acs.jcim.5b00654",

    "[15] Halgren TA. Merck molecular force field. I. Basis, form, scope, parameterization, and "
    "performance of MMFF94. J Comput Chem. 1996;17(5–6):490–519.",

    "[16] Rappé AK, Casewit CJ, Colwell KS, Goddard WA, Skiff WM. UFF, a full periodic table "
    "force field for molecular mechanics and molecular dynamics simulations. J Am Chem Soc. "
    "1992;114(25):10024–10035.",

    "[17] Klamt A, Schuurmann G. COSMO: a new approach to dielectric screening in solvents with "
    "explicit expressions for the screening energy and its gradient. J Chem Soc Perkin Trans 2. "
    "1993;(5):799–805.",
]

for ref in references:
    doc.add_paragraph(ref)

doc.add_paragraph()

# ── Action required summary ────────────────────────────────────────────────
p = doc.add_paragraph()
p.add_run("Summary of [ACTION REQUIRED] items for authors:").bold = True

actions = [
    "1. Paste the merged Introduction at the marked placeholder and verify all in-text citation "
    "numbers align with the reference list above.",
    "2. Fill all [PLACEHOLDER] tags in Results and Discussion with real benchmarking data before "
    "submission.",
    "3. Confirm the Zenodo DOI and replace [ZENODO_DOI_PLACEHOLDER] throughout.",
    "4. Confirm funding details for the Acknowledgements.",
    "5. Verify the MOPAC licence statement (listed as Apache-2.0; confirm it matches the exact "
    "release being bundled).",
]

for a in actions:
    doc.add_paragraph(a)

doc.save("smiles2docking_article.docx")
print("Done: smiles2docking_article.docx")
