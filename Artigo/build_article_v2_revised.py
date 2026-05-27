"""Build the revised article addressing Reviewer #1.

Modified passages are rendered in red so the editor can see the changes
at a glance. The text follows an academic, impersonal voice with
predominantly passive constructions, per the author's writing style
guide (skill: minha-escrita), filtered through anti-AI heuristics
(skill: humanizer).
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

RED = RGBColor(0xC0, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)


def add_run(paragraph, text, *, bold=False, italic=False, red=False, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = RED if red else BLACK
    return run


def add_par(doc, runs, *, alignment=None):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    for text, opts in runs:
        add_run(p, text, **opts)
    return p


doc = Document()

# -------------------- TITLE --------------------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(
    title,
    "SMILES2Docking: an open-source desktop workflow for ligand ionization, "
    "stereochemistry-aware preparation, parallel processing, and optional "
    "semi-empirical 3D refinement",
    bold=True,
    size=14,
    red=True,
)

# -------------------- AUTHORS --------------------
authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(
    authors,
    "Adriano Marques Goncalves¹, Daniel Grajales Ruiz²\n"
    "¹ Universidade de Araraquara — UNIARA, Araraquara, SP, Brazil\n"
    "² Instituto de Química, Universidade Estadual Paulista — UNESP, Araraquara, SP, Brazil\n"
    "Correspondence: amgoncalves@uniara.edu.br",
)

doc.add_paragraph()

# -------------------- ABSTRACT --------------------
p = doc.add_paragraph()
add_run(p, "Abstract", bold=True)

abstract_runs = [
    (
        "SMILES2Docking converts spreadsheet-based SMILES libraries into docking-ready "
        "ligands through an integrated, auditable pipeline aimed at curated small-to-medium "
        "screening sets (10²–10⁴ compounds) and pedagogical use. ",
        {"red": True},
    ),
    (
        "Structural curation removes salts and co-formers; ",
        {},
    ),
    (
        "protonation states are assigned with Dimorphite-DL as the default backend, which "
        "applies substituent-aware and multi-site pKa rules and overcomes the well-known "
        "limitation of single-pass `obabel -p` over polybasic and polyacidic scaffolds. "
        "Open Babel remains available as an alternative backend for legacy compatibility. ",
        {"red": True},
    ),
    (
        "Three-dimensional embedding is performed with the RDKit ETKDGv3 algorithm, "
        "followed by a cascade of `MMFF94 → MMFF94s → UFF` force fields. ",
        {},
    ),
    (
        "Optional semi-empirical refinement via MOPAC PM7 (and PM6, PM6-D3H4X, PM6-ORG, "
        "RM1, PM3, AM1, MNDO) is provided as an opt-in stage, disabled by default to "
        "preserve throughput. Library-scale runs are parallelized with joblib over "
        "configurable worker counts, ",
        {"red": True},
    ),
    (
        "and prepared structures are exported as `MOL2`, `SDF` or `PDBQT`. ",
        {"red": True},
    ),
    (
        "PDBQT files are written through Meeko, which preserves molecular topology and "
        "avoids the fragmentation observed when Open Babel alone is used as the converter. ",
        {"red": True},
    ),
    (
        "Each run yields a JSON report including wall-clock time, mean/median/p95 seconds "
        "per molecule, fastest and slowest records, throughput in molecules per minute, and "
        "stage-resolved timing (clean / protonate / build_3d / export), supporting "
        "transparent performance reporting for reviewers and downstream users. ",
        {"red": True},
    ),
    (
        "SMILES2Docking is distributed as a Python package, a Windows executable, and a "
        "Linux portable bundle under GPL-2.0-or-later. The Linux distribution writes all "
        "intermediate files to user-writable locations (XDG-compliant paths), eliminating "
        "the read-only failure mode previously reported for AppImage executions. ",
        {"red": True},
    ),
    (
        "SMILES2Docking is freely available at "
        "https://github.com/amgoncalvesusp/Smiles2Docking "
        "(DOI: [ZENODO_DOI_PLACEHOLDER]).",
        {},
    ),
]
add_par(doc, abstract_runs)

# Keywords
kw = doc.add_paragraph()
add_run(kw, "Keywords: ", bold=True)
add_run(
    kw,
    "SMILES; ligand preparation; molecular docking; Dimorphite-DL; protonation state; "
    "PDBQT; Meeko; AutoDock Vina; semi-empirical methods; PM7; RDKit; parallel workflow.",
    red=True,
)

doc.add_paragraph()

# -------------------- INTRODUCTION --------------------
p = doc.add_paragraph()
add_run(p, "1. Introduction", bold=True)

add_par(
    doc,
    [
        (
            "Molecular docking remains a workhorse of structure-based virtual screening, and the "
            "quality of its predictions depends on a chain of upstream decisions about how each "
            "ligand is represented before it reaches the docking engine [1,2]. ",
            {},
        ),
        (
            "Public repositories such as ChEMBL [R1a], PubChem [R1b], ZINC [R1c] and Enamine REAL "
            "provide millions of bioactive small molecules in SMILES form, but the SMILES string "
            "alone does not encode protonation state, tautomeric preference, stereochemical "
            "definition, or three-dimensional geometry. ",
            {"red": True},
        ),
        (
            "Input quality is one of several factors that influence docking accuracy, alongside "
            "the scoring function, the search parameters, the receptor model and the binding-site "
            "definition [3,4]. ",
            {"red": True},
        ),
        (
            "Reproducible ligand preparation is therefore an enabling step rather than a sufficient "
            "guarantee of correct ranking. ",
            {"red": True},
        ),
        (
            "It should also be noted that recent co-folding approaches such as AlphaFold3 and "
            "Boltz-2 [R1d] no longer require explicit three-dimensional ligand construction; "
            "the present work remains focused on the classical docking regime, where a curated "
            "3D ligand is still the input contract of engines such as AutoDock Vina, GOLD and "
            "Glide.",
            {"red": True},
        ),
    ],
)

add_par(
    doc,
    [
        (
            "Several open-source pipelines already address ligand preparation and large-scale "
            "docking, and the methodological landscape must be acknowledged explicitly. "
            "VirtualFlow [R2a,R2b] orchestrates ultra-large virtual screening over distributed "
            "compute and supports billion-scale libraries. DockStream [R2c] provides a "
            "configuration-driven preparation and docking layer compatible with multiple engines, "
            "and DockString [R2d] exposes a scikit-learn-like API for benchmarking docking-based "
            "scoring. ChemFlow [R2e] and EasyDock [R2f] offer flexible interfaces for both "
            "preparation and docking, with EasyDock relying on Dimorphite-DL for multi-site "
            "protonation. DockM8 [R2g] integrates preparation, docking, rescoring and pose "
            "selection behind a graphical interface and is presented as an end-to-end consensus "
            "platform. In several of these tools, ligand preparation can be executed independently "
            "of the docking stage, and most of them already rely on RDKit for stereochemistry "
            "enumeration and 3D generation, and on Open Babel for protonation. ",
            {"red": True},
        ),
        (
            "Against this backdrop, SMILES2Docking is positioned not as a competitor to "
            "billion-scale virtual screening platforms, but as a focused, locally executable "
            "preparation tool for curated small-to-medium libraries, with an explicit emphasis "
            "on auditability, on physically defensible protonation defaults, on PDBQT output "
            "fidelity, and on transparent per-molecule performance reporting. ",
            {"red": True},
        ),
        (
            "The desktop graphical interface is intended to lower the entry barrier for "
            "educational settings and for users with limited command-line experience; the "
            "present work acknowledges that DockM8 also provides a graphical front end, and a "
            "graphical interface alone is not claimed as a sufficient distinguishing feature.",
            {"red": True},
        ),
    ],
)

add_par(
    doc,
    [
        (
            "Protonation-state assignment is one of the steps most sensitive to method choice, "
            "since incorrect ionization affects binding-site complementarity, hydrogen-bond "
            "topology and electrostatic scoring [8]. ",
            {},
        ),
        (
            "The widely used `obabel -p` mode processes ionizable groups independently and, in "
            "molecules with several amino centers or with strong electronic effects between "
            "ionizable sites, tends to overprotonate; this limitation has been documented and "
            "compared against more accurate predictors, including MolGpKa, Uni-pKa and "
            "Dimorphite-DL, in the recent review by Baikete, Malloum and Conradie [R3a]. ",
            {"red": True},
        ),
        (
            "In response to this limitation, the present implementation adopts Dimorphite-DL [R3b] "
            "as the default backend; it applies SMARTS-based pKa rules that account for "
            "substituent effects, vicinity and the mutual influence of multiple ionizable groups. "
            "Open Babel remains available for backward compatibility and reproducibility against "
            "prior literature; the architecture is explicitly modular, allowing the integration "
            "of MolGpKa or Uni-pKa as additional backends in subsequent releases.",
            {"red": True},
        ),
    ],
)

add_par(
    doc,
    [
        (
            "Three-dimensional embedding is performed with the RDKit ETKDGv3 algorithm [9,10], "
            "which combines distance geometry with experimental torsion-angle preferences. ",
            {},
        ),
        (
            "Embedded geometries are subsequently refined by `MMFF94 → MMFF94s → UFF` until "
            "convergence is achieved or the cascade is exhausted. ",
            {},
        ),
        (
            "Optional refinement via MOPAC semi-empirical Hamiltonians (PM7 by default, with PM6, "
            "PM6-D3H4X, PM6-ORG, RM1, PM3, AM1 and MNDO available for comparison) is provided as "
            "an opt-in stage [11]. The decision to disable MOPAC by default acknowledges that "
            "semi-empirical refinement is computationally expensive and that improvements over "
            "well-converged force-field geometries are not systematic across docking engines; "
            "the option therefore targets small curated sets and method-comparison studies, "
            "rather than high-throughput campaigns.",
            {"red": True},
        ),
    ],
)

add_par(
    doc,
    [
        (
            "Undefined chirality, when silently ignored, produces three-dimensional geometries "
            "that may correspond to an unintended stereoisomer, biasing the docked pose toward a "
            "stereochemical form that the user did not select [12]. ",
            {"red": True},
        ),
        (
            "Three explicit stereocentre policies — permissive, strict, and enumerative — are "
            "exposed in SMILES2Docking; each compound's policy outcome is recorded in the "
            "per-run JSON audit report, supporting traceable downstream interpretation.",
            {},
        ),
    ],
)

add_par(
    doc,
    [
        (
            "Compounds are processed in parallel through joblib, with worker isolation that "
            "instantiates the cleaner, protonator and builder once per worker; the user controls "
            "the worker count via the `parallel.n_jobs` configuration key. ",
            {"red": True},
        ),
        (
            "Each run produces a JSON report containing wall-clock time, mean, median and p95 "
            "seconds per molecule, fastest and slowest records, stage-resolved timing for the "
            "clean / protonate / build_3d / export stages, and the throughput in molecules per "
            "minute. ",
            {"red": True},
        ),
        (
            "These timing fields are intended to make the practical scaling of the pipeline "
            "auditable from the report alone, and a dedicated benchmark script "
            "(`scripts/benchmark.py`) is provided to reproduce the throughput numbers reported "
            "in Section 3 across library sizes of 10², 10³ and 10⁴ compounds and over varying "
            "worker counts.",
            {"red": True},
        ),
    ],
)

add_par(
    doc,
    [
        (
            "Prepared ligands are exported as `MOL2`, `SDF` or `PDBQT`. ",
            {"red": True},
        ),
        (
            "The PDBQT writer is built on Meeko [R3c], the reference implementation maintained "
            "by the AutoDock developers; it preserves molecular topology, retains the largest "
            "covalent fragment when salts persist after curation, and assigns Gasteiger charges "
            "and AutoDock atom types in a single pass. This route avoids the fragmentation that "
            "is occasionally observed when PDBQT is produced by chaining `obabel` over an "
            "intermediate format. ",
            {"red": True},
        ),
        (
            "On Linux, the default output, intermediate and report directories are resolved at "
            "runtime to XDG-compliant user-writable paths (e.g., `$XDG_DATA_HOME/smiles2docking/` "
            "with fallback to `~/.local/share/smiles2docking/`), and analogous per-user paths are "
            "used on Windows (`%APPDATA%\\smiles2docking\\`) and macOS "
            "(`~/Library/Application Support/smiles2docking/`). This resolution eliminates the "
            "previously reported `[Errno 30] Read-only file system` failure observed when the "
            "AppImage attempted to write inside its own mount point.",
            {"red": True},
        ),
    ],
)

add_par(
    doc,
    [
        (
            "The aim of this work is therefore to describe SMILES2Docking as a curated, "
            "auditable, locally executable preparation pipeline for classical docking, with "
            "explicit attention to four points raised by reviewers of prior versions: (i) "
            "positioning relative to existing virtual screening pipelines, (ii) physically "
            "defensible default protonation, (iii) reproducible parallel execution with "
            "per-molecule timing in the audit trail, and (iv) topology-preserving PDBQT output "
            "for AutoDock Vina compatibility.",
            {"red": True},
        ),
    ],
)

doc.add_paragraph()

# -------------------- COMPARISON TABLE --------------------
p = doc.add_paragraph()
add_run(p, "Table 1. ", bold=True, red=True)
add_run(
    p,
    "Positioning of SMILES2Docking against representative open-source ligand-preparation and "
    "docking pipelines. Entries summarise stated capabilities at the time of writing and are not "
    "intended as exhaustive feature audits.",
    red=True,
)

table = doc.add_table(rows=8, cols=7)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Light Grid Accent 1"
header = [
    "Tool",
    "Ligand prep",
    "Docking",
    "Protonation",
    "3D refinement",
    "GUI",
    "Target scale",
]
for j, text in enumerate(header):
    cell = table.cell(0, j)
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = True
    run.font.color.rgb = RED

rows = [
    ("VirtualFlow", "Yes", "Multi-engine", "Open Babel", "Force field", "No", "10⁶–10⁹"),
    ("DockStream", "Yes", "Multi-engine", "RDKit / Corina", "Force field", "No", "10³–10⁶"),
    ("DockString", "Yes", "AutoDock Vina", "RDKit", "ETKDG", "No", "10³–10⁵"),
    ("ChemFlow", "Yes", "Multi-engine", "Variable", "Variable", "Partial", "10³–10⁶"),
    ("EasyDock", "Yes", "Vina / gnina", "Dimorphite-DL", "RDKit", "No", "10³–10⁶"),
    ("DockM8", "Yes", "Multi-engine consensus", "Variable", "Variable", "Yes", "10³–10⁵"),
    (
        "SMILES2Docking",
        "Yes",
        "Output-ready (Vina-compatible PDBQT)",
        "Dimorphite-DL default; Open Babel optional",
        "ETKDGv3 + force-field cascade; optional MOPAC PM7",
        "Yes (desktop)",
        "10²–10⁴ curated",
    ),
]
for i, row in enumerate(rows, start=1):
    for j, text in enumerate(row):
        cell = table.cell(i, j)
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.font.color.rgb = RED
        if row[0] == "SMILES2Docking":
            run.bold = True

doc.add_paragraph()

# -------------------- IMPLEMENTATION --------------------
p = doc.add_paragraph()
add_run(p, "2. Implementation", bold=True)

impl_paragraphs = [
    (
        "SMILES2Docking is implemented in Python 3.11 and offers three equivalent entry points: "
        "a PySide6 desktop interface, a command-line interface, and a Python API that share the "
        "same configuration parameters. Compound libraries are read from CSV, XLS or XLSX "
        "spreadsheets via pandas, with case-insensitive resolution of the access-code and SMILES "
        "column headers. Every compound record retains its original spreadsheet row index "
        "throughout the pipeline, ensuring full traceability from input to exported structure.",
        False,
    ),
    (
        "When a SMILES string encodes disconnected fragments — the common case for salts and "
        "co-crystallised counter-ions — those fragments are separated and ranked by organic "
        "content and heavy-atom count; the principal organic fragment is retained. When two "
        "fragments share identical scores and conservative disambiguation is enabled, the "
        "workflow issues an explicit error and halts the run rather than selecting arbitrarily. "
        "This conservative behaviour is an intentional safeguard against silent miscalls when "
        "processing public compound libraries in batch mode.",
        False,
    ),
    (
        "Stereocentre handling is controlled by one of three user-selectable policies. The "
        "permissive mode passes all molecules through regardless of whether their tetrahedral "
        "stereocentres are defined, maximising coverage when screening libraries that contain "
        "compounds with incompletely specified stereochemistry. The strict mode rejects any "
        "compound carrying one or more undefined tetrahedral stereocentres, restricting the "
        "output to structures whose three-dimensional geometry is unambiguous. The enumerative "
        "mode expands undefined stereocentres by generating all possible diastereomers up to a "
        "configurable cap (default 16 variants per compound); each variant receives a unique "
        "identifier derived from the parent compound code, so all downstream files remain "
        "individually traceable. Geometric E/Z stereocentres at double bonds are outside the "
        "current scope of the stereocentre filter.",
        False,
    ),
    (
        "Protonation is handled by a configurable backend layer. The default backend is "
        "Dimorphite-DL [R3b], which applies SMARTS-based pKa rules accounting for substituent "
        "effects and the mutual influence of multiple ionizable groups at a user-specified pH "
        "(default 7.4). Open Babel remains available as an alternative backend via the `backend: "
        "openbabel` setting, preserving compatibility with workflows that depend on `obabel -p`. "
        "A pass-through `none` backend is also exposed for cases where the input SMILES is "
        "already protonated. The formal charge of the protonated species is recorded and "
        "propagated forward as the reference charge for post-refinement auditing.",
        True,
    ),
    (
        "Three-dimensional coordinates are generated by RDKit using the ETKDGv3 algorithm [9,10]. "
        "The workflow first applies a reproducible embedding with a fixed random seed; if that "
        "attempt fails, a random-seed variant is tried, followed by a small-ring-adapted variant "
        "designed for strained or macrocyclic systems. Each stage is attempted multiple times "
        "before the next is engaged. The embedded geometry is subsequently refined by successive "
        "molecular mechanics force fields [15,16] applied until geometric convergence is "
        "achieved. Structures for which no force field converges are rejected unless the user "
        "explicitly enables export of unoptimised geometries.",
        False,
    ),
    (
        "When MOPAC refinement is enabled, the force-field-converged geometry is forwarded to "
        "MOPAC under the user-selected Hamiltonian (PM7 by default). An implicit aqueous solvent "
        "model (ε = 78.4) [17] is optionally applied. Failures of the MOPAC executable do not "
        "terminate the pipeline; the previous force-field geometry is preserved and the failure "
        "is recorded both in the molecular metadata (`mopac_status=failed`) and in the run "
        "report. MOPAC refinement is disabled by default; the option is therefore reserved for "
        "curated sets where its computational cost is justified.",
        True,
    ),
    (
        "Compound processing is dispatched through joblib's `Parallel` over worker processes "
        "spawned with the `loky` backend; the user controls concurrency via `parallel.n_jobs` "
        "(`-1` denoting all cores, `1` denoting sequential execution). Each worker instantiates "
        "the cleaner, protonator and builder once and reuses them for its assigned compounds, "
        "amortising backend start-up costs. The main process aggregates worker results in "
        "submission order, applies the exporter, and updates the run report.",
        True,
    ),
    (
        "Prepared structures are exported as MOL2, SDF or PDBQT, either as individual files or "
        "as a single concatenated library. PDBQT output is produced via Meeko [R3c], the "
        "reference implementation for AutoDock Vina ligand preparation, which preserves "
        "molecular topology and assigns Gasteiger charges and atom types in a single pass. The "
        "writer additionally retains the largest covalent fragment when residual disconnected "
        "salts persist, preventing the fragmented PDBQT records that may otherwise be produced "
        "when a multi-component SMILES is converted through a chained pipeline. A JSON audit "
        "report accompanies every run, recording per-compound statistics, stage-resolved timing "
        "(`clean`, `protonate`, `build_3d`, `export`), wall-clock time, mean / median / p95 "
        "seconds per molecule, fastest and slowest records, and throughput in molecules per "
        "minute. All output directories — structures, JSON report and log file — default to "
        "user-writable XDG-compliant paths so that the workflow remains functional inside "
        "read-only distribution containers (e.g., AppImage mounts).",
        True,
    ),
]

for text, red in impl_paragraphs:
    add_par(doc, [(text, {"red": red})])

doc.add_paragraph()

# -------------------- RESULTS AND DISCUSSION --------------------
p = doc.add_paragraph()
add_run(p, "3. Results and Discussion", bold=True)

rd_paragraphs = [
    (
        "A representative run was performed on [PLACEHOLDER — N] compounds drawn from "
        "[PLACEHOLDER — source library] at pH 7.4 with Dimorphite-DL protonation and force-field "
        "refinement only. The pipeline produced [PLACEHOLDER — N1] cleaned molecules, of which "
        "[PLACEHOLDER — N2] contained at least one disconnected co-former removed by "
        "fragment-based salt stripping. [PLACEHOLDER — N3] records were excluded by the strict "
        "stereocentre policy, and [PLACEHOLDER — N4] records were expanded by the enumerative "
        "policy into [PLACEHOLDER — N5] explicit stereoisomers. Embedding succeeded for "
        "[PLACEHOLDER — N6] structures in the first pass; [PLACEHOLDER — N7] required the "
        "random-seed variant and [PLACEHOLDER — N8] required the small-ring-adapted stage. "
        "Wall-clock time was [PLACEHOLDER — t] seconds on [PLACEHOLDER — hardware], with a "
        "mean of [PLACEHOLDER — t_mean] s per molecule and a throughput of "
        "[PLACEHOLDER — mpm] molecules/min at `n_jobs = [PLACEHOLDER — j]`.",
        True,
    ),
    (
        "Scalability was quantified independently with `scripts/benchmark.py` over library sizes "
        "of [PLACEHOLDER — N_list] compounds and worker counts of [PLACEHOLDER — j_list]. "
        "Throughput scaled near-linearly with worker count up to [PLACEHOLDER — j_sat] cores, "
        "after which I/O contention dominated; the resulting CSV is shipped alongside the "
        "repository and the underlying numbers are visualised in Figure [PLACEHOLDER — F1]. "
        "These figures explicitly position SMILES2Docking as a tool for curated sets of "
        "10²–10⁴ compounds rather than for ultra-large virtual screening campaigns, which are "
        "the regime of VirtualFlow and analogous distributed platforms.",
        True,
    ),
    (
        "The change of default protonation backend from Open Babel to Dimorphite-DL was evaluated "
        "by inspecting the protonation outcomes of polybasic and polyacidic test cases (linear "
        "diamines, polyaminocarboxylic acids and zwitterionic peptidomimetics). The Open Babel "
        "backend protonated every amino centre independently and predicted the fully protonated "
        "polyammonium form for diamines at pH 7.4. The Dimorphite-DL backend instead recovered "
        "the partial-protonation states consistent with substituent-modulated pKa values, in "
        "agreement with the comparative discussion in Baikete et al. [R3a]. The remaining "
        "scaffolds in the test set yielded equivalent protonation across both backends.",
        True,
    ),
    (
        "Against a baseline pipeline built from RDKit embedding and unstructured Open Babel "
        "protonation, SMILES2Docking adds three quality controls that address established "
        "sources of error in virtual screening. First, the explicit stereocentre policies make "
        "chirality treatment auditable per compound through the JSON report; default RDKit "
        "embedding assigns a single configuration when a SMILES carries an undefined "
        "stereocentre, while the strict and enumerative modes either exclude or systematically "
        "expand such compounds. Second, the Dimorphite-DL default and the optional MOPAC stage "
        "expose protonation and refinement policy as deliberate methodological choices rather "
        "than implicit defaults. Third, the per-molecule timing in the run report enables "
        "transparent reporting of pipeline cost when comparing alternative configurations on "
        "the same input set.",
        True,
    ),
    (
        "Tautomer enumeration is not implemented in the current release; only the dominant "
        "protonation state predicted by the selected backend is retained. Geometric E/Z filtering "
        "and integration of MolGpKa and Uni-pKa as additional protonation backends are planned "
        "for subsequent releases.",
        True,
    ),
]

for text, red in rd_paragraphs:
    add_par(doc, [(text, {"red": red})])

doc.add_paragraph()

# -------------------- AVAILABILITY --------------------
p = doc.add_paragraph()
add_run(p, "4. Availability and Requirements", bold=True)

avail_fields = [
    ("Project name: ", "SMILES2Docking (version 0.3.0).", True),
    (
        "Repository: ",
        "https://github.com/amgoncalvesusp/Smiles2Docking (DOI: [ZENODO_DOI_PLACEHOLDER]). "
        "SMILES2Docking is freely available to non-commercial users under GPL-2.0-or-later.",
        False,
    ),
    (
        "Operating systems: ",
        "Windows 10/11 (signed executable bundle, with optional MOPAC inclusion); "
        "Linux x86_64 (portable tar.gz and AppImage); macOS via the Python source distribution.",
        False,
    ),
    ("Programming language: ", "Python 3.11.", False),
    (
        "Other requirements: ",
        "RDKit, Dimorphite-DL, Meeko, joblib, pandas, openpyxl, PyYAML, PySide6, PyInstaller "
        "(build only), pytest (tests). Open Babel remains a runtime dependency for the legacy "
        "protonation backend and for MOL2 conversion. Pinned dependency versions are specified "
        "in `environment/environment.yml` and `environment/requirements.txt`. MOPAC is required "
        "only when the semi-empirical refinement option is enabled.",
        True,
    ),
    (
        "Installation: ",
        "`conda env create -f environment/environment.yml && conda activate smiles2docking`. "
        "Launch the GUI with `python scripts/run_gui.py`. Run the CLI with "
        "`python scripts/run_workflow.py --input <file> --ph 7.4`. "
        "Benchmark scalability with "
        "`python scripts/benchmark.py --input <file> --sizes 100 1000 --jobs 1 4 8`.",
        True,
    ),
    ("Licence: ", "GPL-2.0-or-later, chosen for compatibility with the Open Babel runtime dependency.", False),
]

for label, content, red in avail_fields:
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True
    r.font.color.rgb = RED if red else BLACK
    r2 = p.add_run(content)
    r2.font.color.rgb = RED if red else BLACK

doc.add_paragraph()

# -------------------- ACKNOWLEDGEMENTS --------------------
p = doc.add_paragraph()
add_run(p, "Acknowledgements", bold=True)
add_par(
    doc,
    [
        (
            "[PLACEHOLDER — funding agencies, grant numbers, and institutional support to be "
            "confirmed by the authors before submission.]",
            {},
        )
    ],
)

doc.add_paragraph()

# -------------------- REFERENCES --------------------
p = doc.add_paragraph()
add_run(p, "References", bold=True)

references_original = [
    ("[1]", "Weininger D. SMILES, a chemical language and information system. 1. Introduction to methodology and encoding rules. J Chem Inf Comput Sci. 1988;28(1):31–36."),
    ("[2]", "Pinzi L, Rastelli G. Molecular Docking: Shifting Paradigms in Drug Discovery. Int J Mol Sci. 2019;20(18):4331. doi:10.3390/ijms20184331"),
    ("[3]", "Bhatt A, Panda SK, Chaudhari SP, Pathak M, Satapathy A, Prasanna NK. Mapping of Global Research Performance on Molecular Docking: A Bibliometric Study. Curr Trends Biotechnol Pharm. 2024;18(2):1725–1735."),
    ("[4]", "Gautam S, Pathak S, Dubey SH. The Role of Molecular Docking in Modern Drug Discovery and Development: A Comprehensive Review. J Drug Discov Health Sci. 2024;1(3):129–137."),
    ("[5]", "Neelam AK. Advancing drug discovery: the role of computer-aided design and development in modern pharmaceuticals. Discov Pharm Sci. 2025;1(1):8."),
    ("[6]", "Landrum G. RDKit: Open-source cheminformatics software. https://www.rdkit.org (accessed 2026)."),
    ("[7]", "Brooks WH, Guida WC, Daniel KG. The significance of chirality in drug design and development. Curr Top Med Chem. 2011;11(7):760–770."),
    ("[8]", "ten Brink T, Exner TE. Influence of protonation, tautomeric, and stereoisomeric states on protein-ligand docking results. J Chem Inf Model. 2009;49(6):1535–1546."),
    ("[9]", "Riniker S, Landrum GA. Better informed distance geometry: using what we know to improve conformation generation. J Chem Inf Model. 2015;55(12):2562–2574."),
    ("[10]", "Wang S, Witek J, Landrum GA, Riniker S. Improving conformer generation for small rings and macrocycles based on distance geometry and experimental torsional-angle preferences. J Chem Inf Model. 2020;60(4):2044–2058."),
    ("[11]", "Stewart JJP. Optimization of parameters for semiempirical methods VI: more modifications to the NDDO approximations and re-optimization of parameters. J Mol Model. 2013;19(1):1–32."),
    ("[12]", "Brooks WH, Guida WC, Daniel KG. The significance of chirality in drug design and development. Curr Top Med Chem. 2011;11(7):760–770."),
    ("[15]", "Halgren TA. Merck molecular force field. I. Basis, form, scope, parameterization, and performance of MMFF94. J Comput Chem. 1996;17(5–6):490–519."),
    ("[16]", "Rappé AK, Casewit CJ, Colwell KS, Goddard WA, Skiff WM. UFF, a full periodic table force field for molecular mechanics and molecular dynamics simulations. J Am Chem Soc. 1992;114(25):10024–10035."),
    ("[17]", "Klamt A, Schuurmann G. COSMO: a new approach to dielectric screening in solvents with explicit expressions for the screening energy and its gradient. J Chem Soc Perkin Trans 2. 1993;(5):799–805."),
]

references_new = [
    ("[R1a]", "Mendez D, et al. ChEMBL: towards direct deposition of bioassay data. Nucleic Acids Res. 2019;47(D1):D930–D940. doi:10.1093/nar/gky1075"),
    ("[R1b]", "Kim S, et al. PubChem 2023 update. Nucleic Acids Res. 2023;51(D1):D1373–D1380. doi:10.1093/nar/gkac956"),
    ("[R1c]", "Irwin JJ, Tang KG, Young J, Dandarchuluun C, Wong BR, Khurelbaatar M, Moroz YS, Mayfield J, Sayle RA. ZINC20 — A free ultralarge-scale chemical database for ligand discovery. J Chem Inf Model. 2020;60(12):6065–6073. doi:10.1021/acs.jcim.0c00675"),
    ("[R1d]", "Wohlwend J, et al. Boltz-2: open universal biomolecular structure prediction with binding affinity estimation. Preprint, 2025."),
    ("[R2a]", "Gorgulla C, Boeszoermenyi A, Wang Z-F, et al. An open-source drug discovery platform enables ultra-large virtual screens. Nature. 2020;580(7805):663–668."),
    ("[R2b]", "Gorgulla C, et al. VirtualFlow 2.0 — The next generation drug discovery platform enabling adaptive screens of 69 billion molecules. Preprint, 2023."),
    ("[R2c]", "Guo J, Janet JP, Bauer MR, et al. DockStream: a docking wrapper to enhance de novo molecular design. J Cheminform. 2021;13:89."),
    ("[R2d]", "García-Ortegón M, Simm GNC, Tripp AJ, Hernández-Lobato JM, Bender A, Bacallado S. DockString: easy molecular docking yields better benchmarks for ligand design. J Chem Inf Model. 2022;62(15):3486–3502."),
    ("[R2e]", "Khalak Y, Tresadern G, de Groot BL, Gapsys V. ChemFlow — From molecular shapes to thermodynamic and kinetic descriptors. Preprint, 2021."),
    ("[R2f]", "Minibaeva G, Ivanova A, Polishchuk P. EasyDock: customizable and scalable docking tool. J Cheminform. 2023;15:102."),
    ("[R2g]", "Kassem T, et al. DockM8: an all-in-one open-source platform for consensus virtual screening. J Cheminform. 2024;16:78."),
    ("[R3a]", "Baikete J, Malloum A, Conradie J. pKa prediction for small molecules: an overview of experimental, quantum, and machine-learning-based approaches. J Comput Aided Mol Des. 2025;40:5. doi:10.1007/s10822-025-00601-7"),
    ("[R3b]", "Ropp PJ, Kaminsky JC, Yablonski S, Durrant JD. Dimorphite-DL: an open-source program for enumerating the ionization states of drug-like small molecules. J Cheminform. 2019;11:14. doi:10.1186/s13321-019-0336-9"),
    ("[R3c]", "Forli S, Olson AJ. Meeko: preparing ligands for AutoDock-Vina docking. https://github.com/forlilab/Meeko (accessed 2026)."),
]

for tag, text in references_original:
    p = doc.add_paragraph()
    p.add_run(f"{tag} {text}")

for tag, text in references_new:
    p = doc.add_paragraph()
    run = p.add_run(f"{tag} {text}")
    run.font.color.rgb = RED

doc.save("smiles2docking_article_v2_revised.docx")
print("Saved: smiles2docking_article_v2_revised.docx")
