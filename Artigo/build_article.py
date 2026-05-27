from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run(
    "SMILES2DockingFULL: an open-source desktop and command-line workflow "
    "for stereochemistry-aware ligand preparation with PM7 semi-empirical refinement"
)
run.bold = True
run.font.size = Pt(14)

# Authors
authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors.add_run(
    "Adriano Marques Goncalves¹, Daniel Grajales Ruiz²\n"
    "¹ Universidade de Araraquara — UNIARA, Araraquara, SP, Brazil\n"
    "² Instituto de Química, Universidade Estadual Paulista — UNESP, Araraquara, SP, Brazil\n"
    "Correspondence: amgoncalves@uniara.edu.br"
)

doc.add_paragraph()

# Abstract
p = doc.add_paragraph()
p.add_run("Abstract").bold = True

doc.add_paragraph(
    "SMILES2DockingFULL converts spreadsheet-based SMILES libraries into docking-ready ligands through "
    "a single integrated pipeline. It combines salt and co-former removal, a three-mode stereochemistry "
    "policy, pH-dependent protonation via Open Babel, distance-geometry 3D embedding with RDKit, "
    "force-field geometry optimization, and optional PM7 semi-empirical refinement under implicit "
    "solvation. Three stereochemistry modes—permissive, strict, and enumerative—accommodate "
    "both broad library screening and rigorous precision docking. A per-compound JSON audit report "
    "captures stereochemistry actions, protonation states, and post-refinement charge consistency. "
    "The tool ships as a Python package, a Windows executable with bundled MOPAC for zero-configuration "
    "refinement, and a Linux portable bundle under GPL-2.0-or-later. SMILES2DockingFULL is freely "
    "available to non-commercial users at https://github.com/amgoncalvesusp/Smiles2Docking "
    "(DOI: [ZENODO_DOI_PLACEHOLDER])."
)

# Keywords
kw = doc.add_paragraph()
kw.add_run("Keywords: ").bold = True
kw.add_run(
    "SMILES; ligand preparation; molecular docking; stereochemistry enumeration; protonation state; "
    "semi-empirical methods; PM7; RDKit; Open Babel; MOPAC."
)

doc.add_paragraph()

# Introduction placeholder
p = doc.add_paragraph()
p.add_run("[NOTE TO AUTHORS: paste the merged Introduction here]").italic = True

doc.add_paragraph()

# ---- IMPLEMENTATION ----
impl_paragraphs = [
    (
        "SMILES2DockingFULL is implemented in Python 3.11 and offers two equivalent entry points: a PySide6 "
        "desktop interface and a command-line interface that share the same configuration parameters. Compound "
        "libraries are read from CSV, XLS or XLSX spreadsheets via pandas, with case-insensitive resolution of "
        "the access-code and SMILES column headers. Every compound record retains its original spreadsheet row "
        "index throughout the pipeline, ensuring full traceability from input to exported structure."
    ),
    (
        "When a SMILES string encodes disconnected fragments—the common case for salts and co-crystallised "
        "counter-ions—those fragments are separated and ranked by organic content and heavy-atom count; the "
        "principal organic fragment is retained. When two fragments share identical scores and conservative "
        "disambiguation is enabled, the workflow issues an explicit error and halts the run rather than selecting "
        "arbitrarily. This conservative behaviour is an intentional safeguard against silent miscalls when "
        "processing large public compound libraries in batch mode."
    ),
    (
        "Stereocentre handling is controlled by one of three user-selectable policies. The permissive mode passes "
        "all molecules through regardless of whether their tetrahedral stereocentres are defined, maximising "
        "coverage when screening large databases that contain many compounds with incompletely specified "
        "stereochemistry. The strict mode rejects any compound carrying one or more undefined tetrahedral "
        "stereocentres, restricting the output to structures whose three-dimensional geometry is unambiguous. "
        "The enumerative mode expands undefined stereocentres by generating all possible diastereomers up to a "
        "configurable cap (default 16 variants per compound); each variant receives a unique identifier derived "
        "from the parent compound code, so all downstream files remain individually traceable. Geometric E/Z "
        "stereocentres at double bonds are outside the current scope of the stereocentre filter."
    ),
    (
        "Protonation is handled by Open Babel’s pH-dependent protonation module at a user-specified pH "
        "(default 7.4). If the protonated SMILES cannot be directly re-parsed by RDKit, the workflow falls back "
        "to an SDF intermediate before re-importing the structure. The formal charge of the protonated species is "
        "recorded and propagated forward as the reference charge for post-refinement auditing."
    ),
    (
        "Three-dimensional coordinates are generated by RDKit using a distance-geometry-based embedding "
        "algorithm [14]. The workflow first applies a reproducible embedding with a fixed random seed; if that "
        "attempt fails, a random-seed variant is tried, followed by a small-ring-adapted variant designed for "
        "strained or macrocyclic systems. Each stage is attempted multiple times before the next is engaged. "
        "The embedded geometry is subsequently refined by successive molecular mechanics force fields [15,16] "
        "applied until geometric convergence is achieved. Structures for which no force field converges are "
        "rejected unless the user explicitly enables export of unoptimised geometries."
    ),
    (
        "When PM7 semi-empirical refinement is enabled, the force-field-converged geometry is passed to MOPAC "
        "with a standard amide and peptide-bond geometry correction. An implicit aqueous solvent model "
        "(ε = 78.4) [17] is optionally applied to better reflect physiological conditions. Seven "
        "additional semi-empirical Hamiltonians—PM6, PM6-D3H4X, PM6-ORG, RM1, PM3, AM1 and "
        "MNDO—are available for method-comparison studies. The heat of formation from the optimised "
        "geometry is stored as a molecular property. After refinement, a charge audit compares the formal charge "
        "of the MOPAC-optimised structure against the value recorded after protonation. Any discrepancy triggers "
        "an automated charge-rescue heuristic; if that heuristic fails, the record is rejected explicitly, "
        "preventing ligands with silently altered protonation states from reaching the docking engine. On "
        "Windows, MOPAC is bundled within the application executable, so no separate installation is required "
        "to access PM7 refinement on that platform."
    ),
    (
        "Prepared structures are exported in MOL2 or SDF format, as individual files or as single concatenated "
        "libraries, according to user preference. A JSON audit report accompanies every run, recording "
        "per-compound statistics that cover co-former removal actions, protonation state, stereocentre policy "
        "outcomes, embedding and optimisation results, and semi-empirical refinement metadata. This report is "
        "designed to support reproducible reporting and to expose the complete processing history of every "
        "ligand to downstream scrutiny."
    ),
]

for p_text in impl_paragraphs:
    doc.add_paragraph(p_text)

doc.add_paragraph()

# ---- RESULTS AND DISCUSSION ----
rd_paragraphs = [
    (
        "A representative run was performed on [PLACEHOLDER — N] compounds drawn from "
        "[PLACEHOLDER — source library] at pH 7.4 with PM7 refinement and implicit aqueous "
        "solvation. The pipeline produced [PLACEHOLDER — N1] cleaned molecules, of which "
        "[PLACEHOLDER — N2] contained at least one disconnected co-former removed by fragment-based "
        "salt stripping. [PLACEHOLDER — N3] records were excluded by the strict stereocentre policy, "
        "and [PLACEHOLDER — N4] records were expanded by the enumerative policy into "
        "[PLACEHOLDER — N5] explicit stereoisomers, each carrying a unique traceable identifier. "
        "Embedding succeeded for [PLACEHOLDER — N6] structures in the first pass; "
        "[PLACEHOLDER — N7] required the random-seed variant and [PLACEHOLDER — N8] required "
        "the small-ring-adapted stage. PM7 refinement completed for [PLACEHOLDER — N9] structures, "
        "with a mean wall time of [PLACEHOLDER — t] seconds per ligand on "
        "[PLACEHOLDER — hardware]. The mean difference in heat of formation between the force-field "
        "minimum and the PM7 minimum was [PLACEHOLDER — value] kcal/mol."
    ),
    (
        "Against a baseline pipeline built from RDKit embedding and Open Babel protonation alone, "
        "SMILES2DockingFULL adds two quality controls that address established sources of error in virtual "
        "screening. The explicit stereocentre policies make chirality treatment auditable per compound through "
        "the JSON report. Default RDKit embedding assigns an arbitrary configuration when a SMILES carries an "
        "undefined stereocentre; the strict and enumerative modes instead either exclude or systematically "
        "expand such compounds, directly addressing the well-documented consequences of chirality errors on "
        "binding affinity and selectivity [7]. The post-refinement charge audit detects protonation-state "
        "regressions introduced during semi-empirical optimization and either rescues the structure or rejects "
        "it explicitly—a safeguard against the systematic scoring degradation that follows from incorrect "
        "protonation states [8]."
    ),
    (
        "SMILES2DockingFULL also lowers the practical barriers that prevent non-specialist researchers from "
        "conducting well-controlled ligand preparation at scale. Public databases such as ChEMBL, ZINC and "
        "PubChem contain millions of compounds described only as SMILES strings, yet assembling a "
        "production-grade preparation pipeline from individual cheminformatics libraries demands programming "
        "expertise that many end users lack. The three stereocentre modes map directly onto distinct "
        "real-world screening contexts: the permissive mode maximises database coverage for broad hit-finding "
        "campaigns; the strict mode enforces structural rigour for precision docking studies; the enumerative "
        "mode supports systematic exploration of stereoisomeric space for a given scaffold. Bundling MOPAC "
        "within the Windows executable eliminates the most common installation obstacle on that platform, "
        "giving non-expert users access to semi-empirical geometry refinement without any command-line "
        "configuration."
    ),
    (
        "Tautomer enumeration is not currently implemented; only the dominant protonation state predicted by "
        "Open Babel is retained. Compounds are processed sequentially, so wall time scales approximately "
        "linearly with library size. Extension of the stereocentre filter to geometric E/Z stereocentres at "
        "double bonds is planned for a future release."
    ),
]

for p_text in rd_paragraphs:
    doc.add_paragraph(p_text)

doc.add_paragraph()

# ---- AVAILABILITY AND REQUIREMENTS ----
p = doc.add_paragraph()
p.add_run("Availability and Requirements").bold = True

avail_fields = [
    ("Project name: ", "SMILES2DockingFULL (version 0.2.0)."),
    (
        "Repository: ",
        "https://github.com/amgoncalvesusp/Smiles2Docking (DOI: [ZENODO_DOI_PLACEHOLDER]). "
        "SMILES2DockingFULL is freely available to non-commercial users."
    ),
    (
        "Operating systems: ",
        "Windows 10/11 (signed executable bundle; the Windows executable bundles MOPAC, enabling "
        "zero-configuration PM7 refinement on that platform); Linux x86_64 (portable tar.gz and optional "
        "AppImage); macOS via the Python source distribution."
    ),
    ("Programming language: ", "Python 3.11."),
    (
        "Other requirements: ",
        "RDKit, Open Babel, pandas, openpyxl, PyYAML, PySide6, PyInstaller (build only), pytest (tests). "
        "Pinned dependency versions are specified in the repository’s environment file. MOPAC "
        "(Apache-2.0 licence for current open-source releases) is required for PM7 refinement; it is bundled "
        "in the Windows executable and must be installed separately on Linux and macOS."
    ),
    (
        "Installation: ",
        "conda env create -f environment/environment.yml && conda activate smiles2docking. "
        "Launch the GUI with python scripts/run_gui.py or the CLI with "
        "python scripts/run_workflow.py --input <file> --ph 7.4 --pm7 --pm7-solvent --pm7-eps 78.39."
    ),
    ("Licence: ", "GPL-2.0-or-later, chosen for compatibility with the Open Babel runtime dependency."),
]

for label, content in avail_fields:
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    p.add_run(content)

doc.add_paragraph()

# ---- ACKNOWLEDGEMENTS ----
p = doc.add_paragraph()
p.add_run("Acknowledgements").bold = True
doc.add_paragraph(
    "[PLACEHOLDER — funding agencies, grant numbers, and institutional support to be confirmed "
    "by the authors before submission.]"
)

doc.add_paragraph()

# ---- ACTION REQUIRED NOTE ----
p = doc.add_paragraph()
p.add_run(
    "[ACTION REQUIRED: the citation numbers below are assigned on the assumption that the merged Introduction "
    "consumes references [1]–[13] in the order established during the previous drafting session "
    "(Weininger [1], Pinzi & Rastelli [2], Bhatt et al. [3], Gautam et al. [4], Neelam [5], Landrum/RDKit [6], "
    "Brooks et al. [7], ten Brink & Exner [8], O’Boyle et al. [9], Hawkins et al. [10], Sadowski & "
    "Gasteiger [11], Azam & Almahmoud [12], Stewart [13]). When the Introduction is pasted in, verify that "
    "every in-text number aligns with its intended reference and renumber the full list if citations were added "
    "or removed.]"
).italic = True

doc.add_paragraph()

# ---- REFERENCES ----
p = doc.add_paragraph()
p.add_run("References").bold = True

references = [
    "[1] Weininger D. SMILES, a chemical language and information system. 1. Introduction to methodology "
    "and encoding rules. J Chem Inf Comput Sci. 1988;28(1):31–36.",

    "[2] Pinzi L, Rastelli G. Molecular Docking: Shifting Paradigms in Drug Discovery. Int J Mol Sci. "
    "2019;20(18):4331. doi:10.3390/ijms20184331",

    "[3] Bhatt A, Panda SK, Chaudhari SP, Pathak M, Satapathy A, Prasanna NK. Mapping of Global Research "
    "Performance on Molecular Docking: A Bibliometric Study. Curr Trends Biotechnol Pharm. "
    "2024;18(2):1725–1735. doi:10.5530/ctbp.2024.2.21",

    "[4] Gautam S, Pathak S, Dubey SH. The Role of Molecular Docking in Modern Drug Discovery and "
    "Development: A Comprehensive Review. J Drug Discov Health Sci. 2024;1(3):129–137. "
    "doi:10.21590/jddhs.01.03.02",

    "[5] Neelam AK. Advancing drug discovery: the role of computer-aided design and development in modern "
    "pharmaceuticals. Discov Pharm Sci. 2025;1(1):8. doi:10.1007/s44395-025-00008-2",

    "[6] Landrum G. RDKit: Open-source cheminformatics software. https://www.rdkit.org (accessed 2026).",

    "[7] Brooks WH, Guida WC, Daniel KG. The significance of chirality in drug design and development. "
    "Curr Top Med Chem. 2011;11(7):760–770.",

    "[8] ten Brink T, Exner TE. Influence of protonation, tautomeric, and stereoisomeric states on "
    "protein-ligand docking results. J Chem Inf Model. 2009;49(6):1535–1546.",

    "[9] O’Boyle NM, Banck M, James CA, Morley C, Vandermeersch T, Hutchison GR. Open Babel: an open "
    "chemical toolbox. J Cheminform. 2011;3:33.",

    "[10] Hawkins PCD, Skillman AG, Warren GL, Ellingson BA, Stahl MT. Conformer generation with OMEGA: "
    "algorithm and validation using high-quality structures from the Protein Data Bank and Cambridge "
    "Structural Database. J Chem Inf Model. 2010;50(4):572–584.",

    "[11] Sadowski J, Gasteiger J. From atoms and bonds to three-dimensional atomic coordinates: automatic "
    "model builders. Chem Rev. 1993;93(7):2567–2581.",

    "[12] Azam F, Almahmoud SA. Open-Source Molecular Docking and AI-Augmented Structure-Based Drug Design: "
    "Current Workflows, Challenges, and Opportunities. Int J Mol Sci. 2026;27(7):3302. "
    "doi:10.3390/ijms27073302",

    "[13] Stewart JJP. Optimization of parameters for semiempirical methods VI: more modifications to the "
    "NDDO approximations and re-optimization of parameters. J Mol Model. 2013;19(1):1–32.",

    "[14] Riniker S, Landrum GA. Better informed distance geometry: using what we know to improve conformation "
    "generation. J Chem Inf Model. 2015;55(12):2562–2574. doi:10.1021/acs.jcim.5b00654",

    "[15] Halgren TA. Merck molecular force field. I. Basis, form, scope, parameterization, and performance "
    "of MMFF94. J Comput Chem. 1996;17(5–6):490–519.",

    "[16] Rappé AK, Casewit CJ, Colwell KS, Goddard WA, Skiff WM. UFF, a full periodic table force field "
    "for molecular mechanics and molecular dynamics simulations. J Am Chem Soc. "
    "1992;114(25):10024–10035.",

    "[17] Klamt A, Schuurmann G. COSMO: a new approach to dielectric screening in solvents with explicit "
    "expressions for the screening energy and its gradient. J Chem Soc Perkin Trans 2. 1993;(5):799–805.",
]

for ref in references:
    doc.add_paragraph(ref)

doc.add_paragraph()

# ---- ACTION REQUIRED SUMMARY ----
p = doc.add_paragraph()
p.add_run("Summary of [ACTION REQUIRED] items for authors:").bold = True

actions = [
    "1. Paste the merged Introduction text at the marked placeholder and verify all in-text citation "
    "numbers align with the reference list above.",
    "2. Fill all [PLACEHOLDER] tags in the Results and Discussion with real benchmarking data before submission.",
    "3. Confirm the Zenodo DOI and replace [ZENODO_DOI_PLACEHOLDER] throughout.",
    "4. Confirm funding details for the Acknowledgements.",
    "5. Verify the MOPAC licence statement (listed as Apache-2.0; confirm this matches the exact release "
    "being bundled).",
]

for a in actions:
    doc.add_paragraph(a)

doc.save("smiles2docking_article.docx")
print("Saved: smiles2docking_article.docx")
